"""Document formatting engine for DocMetaAI."""

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt


class DocumentFormatter:
    """Apply formatting rules and save a cleaned document."""

    STANDARD_FONT = "Calibri"
    STANDARD_SIZE = 12
    STANDARD_MARGIN = Inches(1)

    def fix(self, file_path: str) -> str:
        """Apply formatting rules and return the output document path."""
        path = Path(file_path)
        document = Document(path)

        self._fix_document_sections(document)
        self._clean_paragraphs(document)
        self._normalize_runs(document)
        self._fix_tables(document)
        self._center_images(document)

        output_path = self._build_output_path(path)
        document.save(output_path)
        return str(output_path)

    def _fix_document_sections(self, document: Document) -> None:
        """Set standard margins and spacing for every section."""
        for section in document.sections:
            section.top_margin = self.STANDARD_MARGIN
            section.bottom_margin = self.STANDARD_MARGIN
            section.left_margin = self.STANDARD_MARGIN
            section.right_margin = self.STANDARD_MARGIN
            section.header_distance = Pt(12)
            section.footer_distance = Pt(12)

    def _clean_paragraphs(self, document: Document) -> None:
        """Clean whitespace, remove empty paragraphs, and normalize paragraph spacing."""
        paragraphs_to_remove = []

        for paragraph in list(document.paragraphs):
            text = paragraph.text or ""
            normalized = self._normalize_whitespace(text)

            if not normalized:
                paragraphs_to_remove.append(paragraph)
                continue

            paragraph.text = normalized
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.15
            if paragraph.style and paragraph.style.name.lower().startswith("heading"):
                self._normalize_heading_style(paragraph)

        for paragraph in paragraphs_to_remove:
            self._delete_paragraph(paragraph)

    def _normalize_whitespace(self, text: str) -> str:
        """Reduce repeated whitespace and trim text."""
        return " ".join(text.split())

    def _normalize_heading_style(self, paragraph: Any) -> None:
        """Standardize heading style names and appearance."""
        if paragraph.style.name.lower().startswith("heading 1"):
            paragraph.style = paragraph.part.styles["Heading 1"]
        elif paragraph.style.name.lower().startswith("heading 2"):
            paragraph.style = paragraph.part.styles["Heading 2"]
        else:
            paragraph.style = paragraph.part.styles.get("Heading 1", paragraph.style)

    def _normalize_runs(self, document: Document) -> None:
        """Standardize fonts and sizes for every run in the document."""
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                run.font.name = self.STANDARD_FONT
                run.font.size = Pt(self.STANDARD_SIZE)
                if run.bold:
                    run.bold = False
                if run.italic:
                    run.italic = False

    def _fix_tables(self, document: Document) -> None:
        """Align tables and normalize cell spacing."""
        for table in document.tables:
            table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        paragraph.paragraph_format.space_after = Pt(4)

    def _center_images(self, document: Document) -> None:
        """Align inline image paragraphs to the center."""
        for paragraph in document.paragraphs:
            if paragraph.runs and any(run.element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/picture}pic') for run in paragraph.runs):
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def _build_output_path(self, path: Path) -> Path:
        """Return a new filename for the fixed document."""
        return path.with_name(f"{path.stem}_Fixed{path.suffix}")

    def _delete_paragraph(self, paragraph: Any) -> None:
        """Delete a paragraph element from the document."""
        element = paragraph._element
        element.getparent().remove(element)
        paragraph._p = paragraph._element = None
