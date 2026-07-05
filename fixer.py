"""Document improvement and structural cleanup for DocMetaAI."""

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


class DocumentFixer:
    """Improve document structure, headings, and readability."""

    STANDARD_FONT = "Calibri"
    STANDARD_SIZE = 12

    def improve(self, file_path: str) -> str:
        """Improve a document and return the improved file path."""
        path = Path(file_path)
        document = Document(path)

        self._remove_extra_spaces(document)
        self._standardize_paragraphs(document)
        self._normalize_headings(document)
        self._standardize_runs(document)

        output_path = self._build_output_path(path)
        document.save(output_path)
        return str(output_path)

    def _remove_extra_spaces(self, document: Document) -> None:
        """Remove repeated whitespace and trim paragraphs."""
        for paragraph in document.paragraphs:
            text = paragraph.text or ""
            paragraph.text = " ".join(text.split())

    def _standardize_paragraphs(self, document: Document) -> None:
        """Apply consistent spacing and alignment for paragraphs."""
        for paragraph in document.paragraphs:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.15
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    def _normalize_headings(self, document: Document) -> None:
        """Improve heading recognition and assign heading styles."""
        for paragraph in document.paragraphs:
            text = (paragraph.text or "").strip()
            if not text:
                continue
            style_name = paragraph.style.name.lower() if paragraph.style else ""
            if style_name.startswith("heading"):
                continue
            # Safely resolve heading styles; python-docx Styles doesn't
            # necessarily support a .get method. Use indexing with fallback.
            try:
                if text.isupper() and len(text) < 120:
                    paragraph.style = document.styles["Heading 1"]
                elif text.endswith(":"):
                    paragraph.style = document.styles["Heading 2"]
            except Exception:
                # If expected heading styles are not present, leave style unchanged.
                continue

    def _standardize_runs(self, document: Document) -> None:
        """Ensure all run text follows the standard font and size."""
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                run.font.name = self.STANDARD_FONT
                run.font.size = Pt(self.STANDARD_SIZE)
                run.bold = False
                run.italic = False

    def _build_output_path(self, path: Path) -> Path:
        """Create a new filename for the improved document."""
        return path.with_name(f"{path.stem}_Improved{path.suffix}")
